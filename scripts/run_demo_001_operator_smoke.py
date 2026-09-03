"""Final local operator smoke for the DEMO-001 reviewable demo path.

Requires the local demo server to be running. The smoke reads Product V1 truth,
re-validates all Top-5 employer-origin URLs live, and explicitly invokes one review-
only application generation for rank 1. It never submits or sends an application.
"""
from __future__ import annotations

from base64 import b64decode
import json
from urllib.parse import urlsplit
from urllib.request import Request, urlopen

from docx import Document
from io import BytesIO
from pypdf import PdfReader
from zipfile import ZipFile

from src.job_lifecycle_health import OUTCOME_SEEN_ACTIVE, JobLifecycleHealthRepository, classify_exact_detail, fetch_exact_detail

AGGREGATOR_HOSTS = (
    "arbeitsagentur.de",
    "gute-jobs.de",
    "stepstone.de",
    "indeed.com",
    "linkedin.com",
)
EXPECTED_FILE_KEYS = {"cv_docx", "cv_pdf", "letter_docx", "letter_pdf", "application_zip"}


class DemoOperatorSmokeStop(RuntimeError):
    pass


def _require(condition: bool, message: str) -> None:
    if not condition:
        raise DemoOperatorSmokeStop(message)


def _host(url: object) -> str:
    return (urlsplit(str(url or "")).hostname or "").casefold()


def _is_aggregator(url: object) -> bool:
    host = _host(url)
    return any(host == item or host.endswith("." + item) for item in AGGREGATOR_HOSTS)


def main() -> int:
    base = "http://127.0.0.1:8781"
    with urlopen(f"{base}/api/v1/product-v1", timeout=15) as response:
        product = json.load(response)

    top = list(product.get("top_jobs") or [])
    _require(len(top) == 5, f"expected exactly five Top jobs, got {len(top)}")
    top.sort(key=lambda row: int(row.get("product_rank") or 999))

    print("=== DEMO-001 OPERATOR SMOKE ===")
    print(f"TOP_JOBS={len(top)}")
    health = JobLifecycleHealthRepository()
    for row in top:
        job_id = int(row["silver_job_id"])
        url = row.get("source_url")
        _require(str(row.get("product_readiness_status") or "") == "rankable", f"Top job not rankable: {job_id}")
        _require(bool(url), f"Top job source URL missing: {job_id}")
        _require(not _is_aggregator(url), f"aggregator leaked into Top-5 action URL: {job_id}|{url}")
        target = health.load_target(job_id)
        probe = fetch_exact_detail(str(url))
        classification = classify_exact_detail(target, probe)
        _require(classification.outcome == OUTCOME_SEEN_ACTIVE, f"Top job not live active: {job_id}|{classification.outcome}")
        print(
            "TOP_LIVE="
            f"{row.get('product_rank')}|{job_id}|score={row.get('overall_quality_score')}|"
            f"{_host(probe.final_url)}|{row.get('company_name')}|{row.get('title')}"
        )

    selected = top[0]
    selected_id = int(selected["silver_job_id"])
    body = json.dumps({"action": "generate_review_draft", "silver_job_id": selected_id}).encode("utf-8")
    request = Request(
        f"{base}/api/v1/product-v1/application-draft",
        data=body,
        headers={"Content-Type": "application/json", "Accept": "application/json"},
        method="POST",
    )
    with urlopen(request, timeout=180) as response:
        draft = json.load(response)

    _require(draft.get("status") == "draft_for_review", "application did not reach draft_for_review")
    _require(int(draft.get("database_writes") or 0) == 0, "application generation wrote database state")
    _require(int(draft.get("submission_writes") or 0) == 0, "application generation wrote submission state")
    _require(int(draft.get("send_actions") or 0) == 0, "application generation performed send action")

    package = draft.get("document_package") or {}
    _require(package.get("status") == "ready_for_download", "document package not ready")
    files = package.get("files") or []
    keys = {str(item.get("key")) for item in files}
    _require(keys == EXPECTED_FILE_KEYS, f"unexpected application file keys: {sorted(keys)}")

    by_key = {str(item["key"]): b64decode(str(item["content_base64"])) for item in files}
    _require(by_key["cv_docx"].startswith(b"PK"), "CV DOCX invalid")
    _require(by_key["letter_docx"].startswith(b"PK"), "letter DOCX invalid")
    _require(by_key["cv_pdf"].startswith(b"%PDF"), "CV PDF invalid")
    _require(by_key["letter_pdf"].startswith(b"%PDF"), "letter PDF invalid")
    _require(by_key["application_zip"].startswith(b"PK"), "application ZIP invalid")

    cv_doc = Document(BytesIO(by_key["cv_docx"]))
    letter_doc = Document(BytesIO(by_key["letter_docx"]))
    cv_text = "\n".join(p.text for p in cv_doc.paragraphs)
    letter_text = "\n".join(p.text for p in letter_doc.paragraphs)
    _require("BASISLEBENSLAUF" not in cv_text, "legacy BASISLEBENSLAUF marker still present")
    _require(letter_text.casefold().count("mit freundlichen grüßen") == 1, "letter greeting/signoff duplicated")

    cv_pdf = PdfReader(BytesIO(by_key["cv_pdf"]))
    letter_pdf = PdfReader(BytesIO(by_key["letter_pdf"]))
    _require(len(cv_pdf.pages) >= 1, "CV PDF has no pages")
    _require(len(letter_pdf.pages) >= 1, "letter PDF has no pages")
    with ZipFile(BytesIO(by_key["application_zip"])) as archive:
        names = set(archive.namelist())
        _require("manifest.json" in names, "ZIP manifest missing")
        _require(len(names) == 5, f"ZIP should contain four documents plus manifest, got {len(names)}")

    print(f"APPLICATION_JOB={selected_id}|{selected.get('company_name')}|{selected.get('title')}")
    print(f"DRAFT_MODE={draft.get('draft_mode')}")
    print(f"PROVIDER_REQUESTS={draft.get('provider_requests')}")
    print(f"APPLICATION_FILES={len(files)}")
    print(f"CV_DOCX_PARAGRAPHS={len(cv_doc.paragraphs)}")
    print(f"CV_PDF_PAGES={len(cv_pdf.pages)}")
    print(f"LETTER_DOCX_PARAGRAPHS={len(letter_doc.paragraphs)}")
    print(f"LETTER_PDF_PAGES={len(letter_pdf.pages)}")
    print("DATABASE_WRITES=0")
    print("SUBMISSION_WRITES=0")
    print("SEND_ACTIONS=0")
    print("DEMO_001_OPERATOR_SMOKE=PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
