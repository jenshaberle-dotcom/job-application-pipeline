from io import BytesIO

from docx import Document
from pypdf import PdfReader

from src.search_intelligence.product_v1_application_document_export import (
    ApplicationDocumentTextBundle,
    render_application_document_bundle,
)


def _bundle() -> ApplicationDocumentTextBundle:
    return ApplicationDocumentTextBundle(
        silver_job_id=434,
        company_name="Muster GmbH & Co. KG",
        cv_text=(
            "Jens Haberle\n\n"
            "Profil\n"
            "Datenengineering mit Python, SQL und zuverlässiger Verarbeitung.\n\n"
            "Erfahrung\n"
            "Aufbau einer strukturierten Job- und Unternehmensdatenpipeline."
        ),
        letter_text=(
            "Sehr geehrte Damen und Herren,\n\n"
            "die ausgeschriebene Data-Engineering-Rolle verbindet Datenplattformen "
            "mit zuverlässiger Verarbeitung.\n\n"
            "Mit freundlichen Grüßen\nJens Haberle"
        ),
    )


def test_four_file_export_contains_editable_docx_and_pdf_variants() -> None:
    files = render_application_document_bundle(_bundle())

    assert [item.key for item in files] == [
        "cv_docx",
        "cv_pdf",
        "letter_docx",
        "letter_pdf",
    ]
    assert len({item.filename for item in files}) == 4
    assert all(item.content_sha256 for item in files)
    assert all(len(item.content_sha256) == 64 for item in files)

    by_key = {item.key: item for item in files}
    assert by_key["cv_docx"].content.startswith(b"PK")
    assert by_key["letter_docx"].content.startswith(b"PK")
    assert by_key["cv_pdf"].content.startswith(b"%PDF")
    assert by_key["letter_pdf"].content.startswith(b"%PDF")


def test_rendered_documents_keep_supplied_complete_text() -> None:
    files = {item.key: item for item in render_application_document_bundle(_bundle())}

    cv_doc = Document(BytesIO(files["cv_docx"].content))
    cv_text = "\n".join(paragraph.text for paragraph in cv_doc.paragraphs)
    assert "Datenengineering mit Python, SQL" in cv_text
    assert "Job- und Unternehmensdatenpipeline" in cv_text

    letter_doc = Document(BytesIO(files["letter_docx"].content))
    letter_text = "\n".join(paragraph.text for paragraph in letter_doc.paragraphs)
    assert "Sehr geehrte Damen und Herren" in letter_text
    assert "Mit freundlichen Grüßen" in letter_text

    cv_pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(files["cv_pdf"].content)).pages
    )
    assert "Datenengineering mit Python, SQL" in cv_pdf_text

    letter_pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(files["letter_pdf"].content)).pages
    )
    assert "Data-Engineering-Rolle" in letter_pdf_text


def test_export_manifest_is_safe_and_company_filename_is_normalized() -> None:
    files = render_application_document_bundle(_bundle())

    for item in files:
        manifest = item.manifest_entry()
        assert "content" not in manifest
        assert manifest["byte_count"] == len(item.content)
        assert "muster-gmbh-co-kg" in item.filename
        assert "/" not in item.filename
        assert "\\" not in item.filename
